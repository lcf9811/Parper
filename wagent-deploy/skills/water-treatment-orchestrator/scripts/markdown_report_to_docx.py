# -*- coding: utf-8 -*-
"""
将水处理遥测研判 Markdown 转为 Word（.docx）。
归档路径：{workspace}/水处理/水处理遥测研判报告_YYYYMMDD_HHMMSS.docx

本文件位于 water-treatment-orchestrator/scripts/，供 SKILL §5.6 在 Markdown 定稿后调用；
skill 可能安装在以下位置之一：
  1. 项目仓库：{project_root}/openclaw_skills/water-treatment-orchestrator/
  2. OpenClaw 安装目录：C:/Users/%user%/.openclaw/skills/water-treatment-orchestrator/
  3. OpenClaw Workspace：C:/Users/%user%/.openclaw/workspace/skills/water-treatment-orchestrator/

调用方式：
  - 推荐：通过 --workspace 参数显式指定 workspace 根目录
  - 或设置环境变量 WATER_PROJECT_ROOT 或 OPENCLAW_WORKSPACE
  - 或从项目仓库根目录运行（脚本会自动推断）
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def _find_workspace_root() -> Path | None:
    """
    尝试多种策略查找 workspace 根目录：
    1. 从脚本位置向上查找，寻找包含 "水处理" 目录或特定标记文件的目录
    2. 检查常见路径模式
    返回 None 如果无法确定
    """
    script_dir = Path(__file__).resolve().parent
    
    # 策略1：向上查找，寻找包含 "水处理" 目录或 app.py/README.md 的目录
    for parent in script_dir.parents:
        # 检查是否是 water_project 根目录的特征
        if (parent / "水处理").exists() or (parent / "app.py").exists() or (parent / "README.md").exists():
            return parent
        # 避免无限向上查找，限制在合理范围内
        if len(parent.parts) < 3:
            break
    
    return None


def _workspace_root() -> Path:
    """
    确定 workspace 根目录，按优先级：
    1. 环境变量 WATER_PROJECT_ROOT
    2. 环境变量 OPENCLAW_WORKSPACE  
    3. 自动推断（从脚本位置向上查找）
    4. 当前工作目录
    """
    # 1. 检查环境变量
    for env_var in ["WATER_PROJECT_ROOT", "OPENCLAW_WORKSPACE"]:
        if env_var in os.environ:
            path = Path(os.environ[env_var])
            if path.exists():
                return path.resolve()
    
    # 2. 尝试自动推断
    inferred = _find_workspace_root()
    if inferred:
        return inferred
    
    # 3. 使用当前工作目录
    return Path.cwd().resolve()


def _default_out_path(workspace: Path, ts: str) -> Path:
    base = workspace / "水处理"
    base.mkdir(parents=True, exist_ok=True)
    return base / f"水处理遥测研判报告_{ts}.docx"


def _set_cell_shading(cell, fill_hex: str) -> None:
    """表头浅底纹。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _add_runs_with_bold(paragraph: Paragraph, text: str) -> None:
    """将 **粗体** 拆成多个 run。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inner = s.strip("|").replace(" ", "")
    return bool(re.match(r"^:?-{3,}(:?\|:?-{3,})*:?$", inner))


def _parse_table_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [c.strip() for c in raw.split("|")]


def _flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = tbl.rows[ri].cells[ci]
            val = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs_with_bold(p, val)
            if ri == 0:
                _set_cell_shading(cell, "E7E6E6")
    doc.add_paragraph("")


def markdown_to_docx(md_text: str, doc: Document) -> None:
    """将常见 Markdown 子集写入 Document：标题、表格、列表、段落、围栏代码块。"""
    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    in_fence = False
    fence_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if in_fence:
            if line.strip().startswith("```"):
                code_text = "\n".join(fence_buf)
                p = doc.add_paragraph()
                run = p.add_run(code_text if code_text else "(空)")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                in_fence = False
                fence_buf = []
            else:
                fence_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_fence = True
            fence_buf = []
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # 标题
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # 表格：以 | 开头且下一行或本块为表格
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines: list[str] = []
            j = i
            while j < len(lines):
                ln = lines[j].strip()
                if not ln:
                    break
                if ln.startswith("|"):
                    table_lines.append(lines[j])
                    j += 1
                    continue
                break
            if len(table_lines) >= 2:
                rows_data: list[list[str]] = []
                for tl in table_lines:
                    if _is_table_separator(tl):
                        continue
                    rows_data.append(_parse_table_row(tl))
                if rows_data:
                    _flush_table(doc, rows_data)
                    i = j
                    continue

        # 列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, stripped[2:].strip())
            i += 1
            continue

        # 普通段落（合并连续非空非特殊行）
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith("```") or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith(
                "- "
            ) or nxt.startswith("* "):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        full = " ".join(para_lines)
        _add_runs_with_bold(p, full)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Markdown 报告 -> Word，输出到 workspace/水处理/水处理遥测研判报告_<时间戳>.docx"
    )
    parser.add_argument(
        "--workspace",
        type=Path,
        default=None,
        help="工作区根目录（默认：从环境变量 WATER_PROJECT_ROOT/OPENCLAW_WORKSPACE 推断，或从脚本位置自动查找）",
    )
    parser.add_argument(
        "--markdown-file",
        type=Path,
        default=None,
        help="Markdown 文件路径（与 stdin 二选一）",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=None,
        help="输出 .docx 完整路径（默认按时间戳自动生成）",
    )
    parser.add_argument(
        "--timestamp",
        default=None,
        help="覆盖时间戳 YYYYMMDD_HHMMSS（默认当前本地时间）",
    )
    args = parser.parse_args()

    workspace = (args.workspace or _workspace_root()).resolve()
    
    # 验证 workspace 是否有效
    if not workspace.exists():
        print(f"错误：工作区目录不存在：{workspace}", file=sys.stderr)
        print(f"提示：请使用 --workspace 参数指定正确的路径，或设置 WATER_PROJECT_ROOT 环境变量", file=sys.stderr)
        return 2

    if args.markdown_file:
        md_path = args.markdown_file.resolve()
        if not md_path.is_file():
            print(f"错误：找不到 Markdown 文件：{md_path}", file=sys.stderr)
            return 2
        md_text = md_path.read_text(encoding="utf-8")
    else:
        md_text = sys.stdin.read()
        if not md_text.strip():
            print("错误：未提供 --markdown-file 且 stdin 为空", file=sys.stderr)
            return 2

    ts = args.timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = args.out or _default_out_path(workspace, ts)
    out_path = out_path.resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sect = doc.sections[0]
    sect.left_margin = sect.right_margin = sect.top_margin = sect.bottom_margin = Pt(72)

    markdown_to_docx(md_text, doc)
    doc.save(str(out_path))
    print(str(out_path))
    return 0


if __name__ == "__main__":
    sys.exit(main())
